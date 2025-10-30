# qna_parse.py
import re
from pathlib import Path
import json

def load_json(p): return json.load(open(p, "r", encoding="utf-8"))

def extract_qna_labeled(text: str):
    pattern = re.compile(
        r"(?:^\s*\d+\.\s*)?"
        r"(?:\*\*)?\s*Question\s*:\s*(.*?)\s*"
        r"(?:\*\*)?\s*Answer\s*:\s*(.*?)\s*"
        r"(?=(?:\n\s*\n|^\s*\d+\.|(?:\*\*)?\s*Question\s*:|\Z))",
        flags=re.IGNORECASE | re.DOTALL | re.MULTILINE,
    )
    pairs = []
    for q, a in pattern.findall(text or ""):
        q = re.sub(r"\s+", " ", q).strip()
        a = re.sub(r"\s+", " ", a).strip()
        if q and a:
            pairs.append({"question": q, "answer": a})
    return pairs

def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")

def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx is required to read .docx") from e
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)

def read_text_auto(path_like) -> str:
    p = Path(path_like)
    suf = p.suffix.lower()
    if suf == ".txt":
        return _read_txt(p)
    if suf == ".docx":
        return _read_docx(p)
    if suf == ".json":
        raise ValueError("read_text_auto(): JSON is structured; use load_json for .json.")
    return _read_txt(p)

def load_pairs_from_path(path: str):
    p = Path(path)
    ext = p.suffix.lower()

    if ext == ".json":
        with p.open("r", encoding="utf-8-sig") as f:
            obj = json.load(f)

        # Accept both formats:
        # 1) Flat array: [ {question, answer, id?}, ... ]
        # 2) Wrapped: { "papers": [...], "items": [ {question, answer, id?}, ... ] }
        if isinstance(obj, dict) and isinstance(obj.get("items"), list):
            data = obj["items"]
        elif isinstance(obj, list):
            data = obj
        else:
            raise ValueError(
                "Unsupported JSON format. Expected a list of QA objects or "
                'a dict with an "items" list.'
            )

        out = []
        for i, ex in enumerate(data):
            if not isinstance(ex, dict):
                raise ValueError(f"Item {i} is not an object: {type(ex)}")
            q = (ex.get("question") or "").strip()
            a = (ex.get("answer") or "").strip()
            if not q or not a:
                raise ValueError(f"Item {i} missing question/answer")
            out.append({"id": ex.get("id") or f"Q{i+1}", "question": q, "answer": a})
        return out
    
    if ext == ".docx":
        try:
            from docx import Document  # python-docx
        except ImportError as e:
            raise RuntimeError("python-docx is required to read DOCX datasets: pip install python-docx") from e

        doc = Document(str(p))
        # Concatenate paragraphs to preserve order; header like "Papers: ..." may be present—harmless.
        text = "\n".join((para.text or "") for para in doc.paragraphs)

        # Prefer structured parser if available
        items = []
        try:
            # If extract_qna_labeled is defined in your package, use it
            from qna_parse import extract_qna_labeled  # safe if same module/package
            items = extract_qna_labeled(text) or []
        except Exception:
            pass

        # Fallback: labeled "Question:" / "Answer:" blocks
        if not items:
            blocks = re.findall(
                r"Question:\s*(.+?)\n\s*Answer:\s*(.+?)(?=\n\s*Question:|\Z)",
                text, flags=re.S | re.I
            )
            items = [{"question": q.strip(), "answer": a.strip()} for q, a in blocks]

        out = [{"id": f"Q{i+1}", "question": it["question"], "answer": it["answer"]} for i, it in enumerate(items)]
        return out

    raise ValueError(f"Unsupported dataset extension: {ext} (use .json, .jsonl, or .docx)")


def infer_pdfs_from_dataset(dataset_path: str) -> list[str]:
    """
    Minimal fallback:
    - JSON: expect top-level {"papers": ["/abs/a.pdf", "/abs/b.pdf", ...]}
    - DOCX: expect the FIRST non-empty paragraph to start with:
        "Papers: /abs/a.pdf, /abs/b.pdf"
    No extra formats are supported.
    """
    ds = Path(dataset_path)
    ext = ds.suffix.lower()

    if ext == ".json":
        data = load_json(ds)
        papers = data.get("papers") if isinstance(data, dict) else None
        if isinstance(papers, list) and papers:
            return [str(Path(p).expanduser().resolve()) for p in papers if isinstance(p, str) and p.strip()]
        return []

    if ext == ".docx":
        try:
            from docx import Document  # python-docx
        except ImportError as e:
            raise RuntimeError("python-docx required to read 'Papers:' from DOCX. Install via: pip install python-docx") from e
        doc = Document(str(ds))
        for para in doc.paragraphs:
            line = (para.text or "").strip()
            if not line:
                continue
            if line.startswith("Papers:"):
                rest = line[len("Papers:"):].strip()
                if rest:
                    parts = [s.strip() for s in rest.split(",") if s.strip()]
                    return [str(Path(p).expanduser().resolve()) for p in parts]
                break
        return []

    # Unsupported dataset container for fallback
    return []

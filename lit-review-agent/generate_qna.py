import os, json, argparse, re
from openai import OpenAI
from pathlib import Path
from qna_parse import extract_qna_labeled

DEF_PROMPT_GEN = "prompts/prompt_qa_gen.txt"
DEF_PROMPT_GEN_MULTI = "prompts/prompt_qa_gen_multi.txt"

def load_text(p):
    return open(p, "r", encoding="utf-8").read()

def save_docx_raw(raw_text: str, out_path: Path, title: str, papers_abs: list[str]):
    """
    Dump the model's raw output text exactly as-is (line-preserving) into a DOCX.
    The FIRST non-empty paragraph will be:
        Papers: /abs/a.pdf, /abs/b.pdf
    Then a title paragraph, then one empty paragraph, then the raw content.
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise RuntimeError("python-docx is required. Install via: pip install python-docx") from e

    doc = Document()

    # Required header line for downstream evaluation fallback
    papers_line = "Papers: " + ", ".join(papers_abs)
    doc.add_paragraph(papers_line)

    # Title
    h_run = doc.add_paragraph().add_run(title)
    h_run.bold = True
    h_run.font.size = Pt(16)

    # Blank spacer
    doc.add_paragraph("")

    # Raw model output, line-preserving
    for line in (raw_text or "").splitlines():
        doc.add_paragraph(line)

    doc.save(str(out_path))


def generate_qa(client: OpenAI,
                pdf_paths,
                prompt_single_path: str,
                prompt_multi_path: str,
                model: str,
                out_dir: Path,
                id_prefix: str = "Q",
                max_items: int | None = None,
                export_docx: bool = False) -> dict:
    """
    Q&A generator: accepts one or many PDFs.
    - If one PDF: uploads that file and produces <stem>_QA.json[/docx].
    - If multiple PDFs: uploads all, runs multi-paper prompt, produces multi_{N}_papers_QA.json[/docx].
    """
    # Normalize input to list
    if isinstance(pdf_paths, (str, Path)):
        pdf_paths = [str(pdf_paths)]

    # Absolute paths (used in JSON 'papers' and DOCX header)
    papers_abs = [str(Path(p).expanduser().resolve()) for p in pdf_paths]

    n = len(pdf_paths)
    is_multi = n > 1

    # Choose prompt + output naming
    prompt_path = prompt_multi_path if is_multi else prompt_single_path
    prompt = load_text(prompt_path)

    out_dir.mkdir(parents=True, exist_ok=True)
    if is_multi:
        stem = f"multi_{n}_papers"
    else:
        stem = Path(pdf_paths[0]).stem

    json_path = out_dir / f"{stem}_QA.json"
    docx_path = out_dir / f"{stem}_QA.docx"

    # Build content: one input_text + 1..N input_file items
    content = [{
        "type": "input_text",
        "text": "Generate the QA benchmark as instructed." + (" Use ALL provided papers." if is_multi else "")
    }]

    # Upload files
    uploaded_ids = []
    for p in pdf_paths:
        with open(p, "rb") as fh:
            f = client.files.create(file=fh, purpose="assistants")
        uploaded_ids.append(f.id)
        content.append({"type": "input_file", "file_id": f.id})

    # Call model
    resp = client.responses.create(
        model=model,
        instructions=prompt,
        input=[{"role": "user", "content": content}],
    )
    raw_text = (resp.output_text or "").strip()

    # Parse Q/A items
    items = extract_qna_labeled(raw_text) or []
    normalized = [
        {"id": f"{id_prefix}{i+1}", "question": it["question"], "answer": it["answer"]}
        for i, it in enumerate(items) if (max_items is None or i < max_items)
    ]

    # Write JSON in the evaluate.py-expected shape:
    # {
    #   "papers": ["/abs/one.pdf", "/abs/two.pdf"],
    #   "items": [ {id, question, answer}, ... ]
    # }
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump({"papers": papers_abs, "items": normalized}, f, ensure_ascii=False, indent=2)
    print(f"[write] json -> {json_path} (n={len(normalized)})")

    # Optional raw DOCX dump with required 'Papers:' first line
    if export_docx:
        save_docx_raw(raw_text, docx_path, f"{stem} — Model Output (Raw)", papers_abs)
        print(f"[write] docx(raw) -> {docx_path} (chars={len(raw_text)})")

    return {"json": json_path, "docx": docx_path if export_docx else None, "papers": papers_abs}


def main():
    ap = argparse.ArgumentParser(
        description="Generate QA benchmarks from one or more PDFs (JSON, optional DOCX). "
                    "Uses a different prompt automatically when multiple PDFs are provided."
    )
    ap.add_argument("--pdf", nargs="+", required=True,
                    help="One or more PDF paths. If >1, a single multi-paper QA is generated.")
    ap.add_argument("--prompt", default=DEF_PROMPT_GEN,
                    help="Instruction prompt for SINGLE-PDF QA.")
    ap.add_argument("--prompt_multi", default=DEF_PROMPT_GEN_MULTI,
                    help="Instruction prompt for MULTI-PDF (cross-paper) QA.")
    ap.add_argument("--model", default="gpt-4o", help="Model for QA generation.")
    ap.add_argument("--out-dir", default="out", help="Directory to write outputs.")
    ap.add_argument("--id-prefix", default="Q", help="ID prefix for items.")
    ap.add_argument("--max-items", type=int, default=None, help="Cap number of QAs in the output.")
    ap.add_argument("--docx", action="store_true", help="Also write a DOCX alongside the JSON (with 'Papers:' header).")
    args = ap.parse_args()

    client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])
    out_dir = Path(args.out_dir)

    res = generate_qa(
        client=client,
        pdf_paths=args.pdf,
        prompt_single_path=args.prompt,
        prompt_multi_path=args.prompt_multi,
        model=args.model,
        out_dir=out_dir,
        id_prefix=args.id_prefix,
        max_items=args.max_items,
        export_docx=args.docx,
    )

    # Make paths printable
    printable = {k: (str(v) if isinstance(v, Path) else v) for k, v in res.items()}
    print(json.dumps({"generated": [printable]}, indent=2))


if __name__ == "__main__":
    main()
